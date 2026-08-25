import io
import re
import zipfile
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass
class MarkdownBlock:
    block_type: str
    text: str = ""
    level: int = 0
    rows: List[List[str]] = field(default_factory=list)
    language: str = ""


@dataclass
class DocumentGenerationResult:
    docx_bytes: bytes
    block_count: int
    markdown_chars: int
    title: str


class MarkdownDocumentParser:
    """把模型提交的 Markdown 正文解析成 Word 渲染层可消费的块结构。"""

    _heading_re = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
    _unordered_re = re.compile(r"^\s*[-*+]\s+(.+?)\s*$")
    _ordered_re = re.compile(r"^\s*\d+[.)]\s+(.+?)\s*$")
    _quote_re = re.compile(r"^\s*>\s?(.*?)\s*$")
    _fence_re = re.compile(r"^\s*```\s*([A-Za-z0-9_+.-]*)\s*$")

    def parse(self, markdown: str) -> List[MarkdownBlock]:
        source = str(markdown or "").replace("\r\n", "\n").replace("\r", "\n")
        lines = source.split("\n")
        blocks: List[MarkdownBlock] = []
        paragraph_lines: List[str] = []
        cursor = 0

        def flush_paragraph() -> None:
            text = " ".join(line.strip() for line in paragraph_lines if line.strip()).strip()
            paragraph_lines.clear()

            if text:
                blocks.append(MarkdownBlock("paragraph", text=text))

        while cursor < len(lines):
            line = lines[cursor]
            stripped = line.strip()

            fence_match = self._fence_re.match(stripped)

            if fence_match:
                flush_paragraph()
                code_lines, next_cursor = self._read_code_block(lines, cursor + 1)
                blocks.append(MarkdownBlock(
                    "code",
                    text="\n".join(code_lines).rstrip(),
                    language=fence_match.group(1).strip(),
                ))
                cursor = next_cursor
                continue

            table_rows, next_cursor = self._read_table(lines, cursor)

            if table_rows:
                flush_paragraph()
                blocks.append(MarkdownBlock("table", rows=table_rows))
                cursor = next_cursor
                continue

            if not stripped:
                flush_paragraph()
                blocks.append(MarkdownBlock("blank"))
                cursor += 1
                continue

            heading_match = self._heading_re.match(stripped)

            if heading_match:
                flush_paragraph()
                blocks.append(MarkdownBlock(
                    "heading",
                    text=heading_match.group(2).strip(),
                    level=min(len(heading_match.group(1)), 4),
                ))
                cursor += 1
                continue

            unordered_match = self._unordered_re.match(line)

            if unordered_match:
                flush_paragraph()
                blocks.append(MarkdownBlock("bullet", text=unordered_match.group(1).strip()))
                cursor += 1
                continue

            ordered_match = self._ordered_re.match(line)

            if ordered_match:
                flush_paragraph()
                blocks.append(MarkdownBlock("number", text=ordered_match.group(1).strip()))
                cursor += 1
                continue

            quote_match = self._quote_re.match(line)

            if quote_match:
                flush_paragraph()
                blocks.append(MarkdownBlock("quote", text=quote_match.group(1).strip()))
                cursor += 1
                continue

            paragraph_lines.append(line)
            cursor += 1

        flush_paragraph()
        return self._trim_blank_edges(blocks)

    def _read_code_block(self, lines: List[str], cursor: int) -> Tuple[List[str], int]:
        code_lines: List[str] = []

        while cursor < len(lines):
            line = lines[cursor]

            if self._fence_re.match(line.strip()):
                return code_lines, cursor + 1

            code_lines.append(line)
            cursor += 1

        return code_lines, cursor

    def _read_table(self, lines: List[str], cursor: int) -> Tuple[List[List[str]], int]:
        if cursor + 1 >= len(lines):
            return [], cursor

        header = lines[cursor].strip()
        separator = lines[cursor + 1].strip()

        if "|" not in header or not self._is_table_separator(separator):
            return [], cursor

        rows = [self._split_table_row(header)]
        cursor += 2

        while cursor < len(lines):
            line = lines[cursor].strip()

            if "|" not in line:
                break

            rows.append(self._split_table_row(line))
            cursor += 1

        column_count = max(len(row) for row in rows)

        if column_count < 2:
            return [], cursor

        normalized_rows = []

        for row in rows:
            normalized = row + [""] * (column_count - len(row))
            normalized_rows.append(normalized[:column_count])

        return normalized_rows, cursor

    def _is_table_separator(self, line: str) -> bool:
        parts = self._split_table_row(line)

        if len(parts) < 2:
            return False

        for part in parts:
            cell = part.strip()

            if not re.fullmatch(r":?-{3,}:?", cell):
                return False

        return True

    def _split_table_row(self, line: str) -> List[str]:
        text = line.strip()

        if text.startswith("|"):
            text = text[1:]

        if text.endswith("|"):
            text = text[:-1]

        return [cell.strip() for cell in text.split("|")]

    def _trim_blank_edges(self, blocks: List[MarkdownBlock]) -> List[MarkdownBlock]:
        start = 0
        end = len(blocks)

        while start < end and blocks[start].block_type == "blank":
            start += 1

        while end > start and blocks[end - 1].block_type == "blank":
            end -= 1

        return blocks[start:end]


class DocxRenderer:
    """把 MarkdownBlock 渲染成 Word 文档二进制。"""

    default_font_name = "Microsoft YaHei"

    def render(self, blocks: List[MarkdownBlock], title: str = "", doc_options: Optional[Dict[str, Any]] = None) -> bytes:
        options = dict(doc_options or {})
        document = Document()
        self._configure_document(document, options)
        clean_title = str(title or "").strip()

        if clean_title:
            paragraph = document.add_paragraph(clean_title, style="Title")
            self._set_paragraph_font(paragraph)

        for block in blocks:
            self._render_block(document, block)

        output = io.BytesIO()
        document.save(output)
        return output.getvalue()

    def _configure_document(self, document, options: Dict[str, Any]) -> None:
        font_name = str(options.get("font_name") or self.default_font_name).strip()
        normal_size = self._read_float_option(options, "font_size", 10.5)
        line_spacing = self._read_float_option(options, "line_spacing", 1.15)

        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(self._read_float_option(options, "top_margin", 1.0))
        section.bottom_margin = Inches(self._read_float_option(options, "bottom_margin", 1.0))
        section.left_margin = Inches(self._read_float_option(options, "left_margin", 1.25))
        section.right_margin = Inches(self._read_float_option(options, "right_margin", 1.25))

        for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
            style = document.styles[style_name]
            style.font.name = font_name

            if style._element.rPr is not None:
                style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        normal_style = document.styles["Normal"]
        normal_style.font.size = Pt(normal_size)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = line_spacing

    def _render_block(self, document, block: MarkdownBlock) -> None:
        if block.block_type == "blank":
            document.add_paragraph("")
            return

        if block.block_type == "heading":
            paragraph = document.add_heading(level=max(1, min(block.level, 4)))
            self._add_inline_runs(paragraph, block.text)
            self._set_paragraph_font(paragraph)
            return

        if block.block_type == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            self._add_inline_runs(paragraph, block.text)
            return

        if block.block_type == "number":
            paragraph = document.add_paragraph(style="List Number")
            self._add_inline_runs(paragraph, block.text)
            return

        if block.block_type == "quote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            run = paragraph.add_run(block.text)
            run.italic = True
            run.font.color.rgb = RGBColor(89, 89, 89)
            self._set_run_font(run)
            return

        if block.block_type == "code":
            self._add_code_block(document, block.text)
            return

        if block.block_type == "table":
            self._add_table(document, block.rows)
            return

        paragraph = document.add_paragraph()
        self._add_inline_runs(paragraph, block.text)

    def _add_inline_runs(self, paragraph, text: str) -> None:
        token_re = re.compile(
            r"(\[([^\]]+)\]\(([^)]+)\)|`([^`]+)`|\*\*([^*]+)\*\*|\*([^*]+)\*)"
        )
        cursor = 0

        for match in token_re.finditer(str(text or "")):
            if match.start() > cursor:
                self._add_text_run(paragraph, text[cursor:match.start()])

            if match.group(2) is not None:
                self._add_hyperlink(paragraph, match.group(2).strip(), match.group(3).strip())
            elif match.group(4) is not None:
                run = self._add_text_run(paragraph, match.group(4), font_name="Consolas")
                run.font.size = Pt(9.5)
            elif match.group(5) is not None:
                run = self._add_text_run(paragraph, match.group(5))
                run.bold = True
            elif match.group(6) is not None:
                run = self._add_text_run(paragraph, match.group(6))
                run.italic = True

            cursor = match.end()

        if cursor < len(text):
            self._add_text_run(paragraph, text[cursor:])

    def _add_text_run(self, paragraph, text: str, font_name: Optional[str] = None):
        run = paragraph.add_run(str(text or ""))
        self._set_run_font(run, font_name=font_name)
        return run

    def _add_hyperlink(self, paragraph, text: str, url: str) -> None:
        if not text or not url:
            self._add_text_run(paragraph, text or url)
            return

        rel_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel_id)

        run_elem = OxmlElement("w:r")
        run_props = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_props.append(color)
        run_props.append(underline)
        run_elem.append(run_props)

        text_elem = OxmlElement("w:t")
        text_elem.text = text
        run_elem.append(text_elem)
        hyperlink.append(run_elem)
        paragraph._element.append(hyperlink)

    def _add_code_block(self, document, text: str) -> None:
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        tc_pr = cell._element.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F5F5F5")
        shading.set(qn("w:val"), "clear")
        tc_pr.append(shading)

        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(str(text or ""))
        self._set_run_font(run, font_name="Consolas", size=9)

    def _add_table(self, document, rows: List[List[str]]) -> None:
        if not rows:
            return

        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"

        for row_index, row in enumerate(rows):
            for column_index in range(column_count):
                cell = table.rows[row_index].cells[column_index]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                self._add_inline_runs(paragraph, row[column_index] if column_index < len(row) else "")

                if row_index == 0:
                    for run in paragraph.runs:
                        run.bold = True

    def _set_paragraph_font(self, paragraph) -> None:
        for run in paragraph.runs:
            self._set_run_font(run)

    def _set_run_font(self, run, font_name: Optional[str] = None, size: Optional[float] = None) -> None:
        final_font_name = font_name or self.default_font_name
        run.font.name = final_font_name

        if size is not None:
            run.font.size = Pt(size)

        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), final_font_name)

    def _read_float_option(self, options: Dict[str, Any], key: str, default: float) -> float:
        value = options.get(key)

        if value is None:
            return default

        try:
            number = float(value)
        except Exception:
            raise ValueError(f"doc_options.{key} 必须是数字")

        if number <= 0:
            raise ValueError(f"doc_options.{key} 必须大于 0")

        return number


class DocxValidator:
    """校验生成结果是否为可被 Word 识别的 docx 包。"""

    required_members = {
        "[Content_Types].xml",
        "word/document.xml",
        "_rels/.rels",
    }

    def validate(self, docx_bytes: bytes) -> None:
        if not isinstance(docx_bytes, (bytes, bytearray)) or not docx_bytes:
            raise ValueError("DOCX 生成结果为空")

        try:
            with zipfile.ZipFile(io.BytesIO(bytes(docx_bytes)), "r") as archive:
                names = set(archive.namelist())
        except zipfile.BadZipFile:
            raise ValueError("DOCX 生成结果不是有效的 OpenXML 压缩包")

        missing_members = sorted(self.required_members - names)

        if missing_members:
            raise ValueError(f"DOCX 生成结果缺少必要文件: {', '.join(missing_members)}")

        Document(io.BytesIO(bytes(docx_bytes)))


class DocumentGenerationService:
    """文档生成编排服务：解析 Markdown、渲染 Word、校验 docx。"""

    def __init__(
        self,
        parser: Optional[MarkdownDocumentParser] = None,
        renderer: Optional[DocxRenderer] = None,
        validator: Optional[DocxValidator] = None,
    ):
        self.parser = parser or MarkdownDocumentParser()
        self.renderer = renderer or DocxRenderer()
        self.validator = validator or DocxValidator()

    def create_docx(self, markdown: str, title: str = "", doc_options: Optional[Dict[str, Any]] = None) -> DocumentGenerationResult:
        source = str(markdown or "")

        if not source.strip():
            raise ValueError("markdown 不能为空")

        blocks = self.parser.parse(source)
        docx_bytes = self.renderer.render(blocks, title=title, doc_options=doc_options)
        self.validator.validate(docx_bytes)

        return DocumentGenerationResult(
            docx_bytes=docx_bytes,
            block_count=len(blocks),
            markdown_chars=len(source),
            title=str(title or "").strip(),
        )
