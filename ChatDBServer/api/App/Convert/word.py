import base64
import json
import os
import re
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from urllib import parse as urllib_parse
from urllib import request as urllib_request

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


WORD_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_MARKDOWN_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_KNOWLEDGE_IMAGE_ROUTE_RE = re.compile(r"^/api/(?:public/)?knowledge/image/([^/]+)/([^/?#]+)$")
_SAFE_USERNAME_RE = re.compile(r"^[\w.\-]+$", re.UNICODE)
_SAFE_IMAGE_ID_RE = re.compile(r"^kimg_[a-z0-9]{16}$")
_IMAGE_MAX_BYTES = 12 * 1024 * 1024


class _HtmlContentParser(HTMLParser):
    block_tags = {
        "article",
        "blockquote",
        "dd",
        "div",
        "dl",
        "dt",
        "figcaption",
        "figure",
        "footer",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "header",
        "li",
        "main",
        "ol",
        "p",
        "pre",
        "section",
        "table",
        "tbody",
        "td",
        "tfoot",
        "th",
        "thead",
        "tr",
        "ul",
    }

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.tokens: List[Dict[str, Any]] = []

    def handle_starttag(self, tag, attrs):
        normalized_tag = str(tag or "").lower()
        attr_map = {str(name or "").lower(): str(value or "") for name, value in attrs}

        if normalized_tag == "hr":
            self.tokens.append({"type": "horizontal_rule"})
            self.tokens.append({"type": "newline"})
            return

        if normalized_tag == "img":
            self.tokens.append({
                "type": "image",
                "src": attr_map.get("src", ""),
                "alt": attr_map.get("alt", ""),
                "width_px": self._read_size_px(attr_map, "width"),
                "height_px": self._read_size_px(attr_map, "height"),
            })
            self.tokens.append({"type": "newline"})
            return

        if normalized_tag == "br":
            self.tokens.append({"type": "newline"})
            return

        if normalized_tag in self.block_tags:
            self.tokens.append({"type": "newline"})

    def handle_endtag(self, tag):
        normalized_tag = str(tag or "").lower()

        if normalized_tag in self.block_tags:
            self.tokens.append({"type": "newline"})

    def handle_data(self, data):
        text = str(data or "")

        if text:
            self.tokens.append({"type": "text", "text": text})

    def _read_size_px(self, attr_map: Dict[str, str], key: str) -> Optional[float]:
        raw = str(attr_map.get(key) or "").strip()

        if raw:
            match = re.match(r"^(\d+(?:\.\d+)?)", raw)

            if match:
                return float(match.group(1))

        style = str(attr_map.get("style") or "")
        match = re.search(rf"{re.escape(key)}\s*:\s*(\d+(?:\.\d+)?)px", style, flags=re.I)

        if match:
            return float(match.group(1))

        return None


class MarkdownWordConverter:
    """Convert Markdown-like content to a Word document."""

    mimetype = WORD_MIMETYPE
    font_name = "Microsoft YaHei"
    max_image_width_inches = 5.8

    def __init__(self, base_dir: Optional[str] = None):
        self.base_dir = Path(base_dir or Path(__file__).resolve().parents[3]).resolve()

    def build(self, contents: Iterable[str]) -> BytesIO:
        document = Document()
        self._configure_styles(document)

        for index, content in enumerate(contents, start=1):

            if index > 1:
                document.add_page_break()

            self._add_content(document, str(content or ""))

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return output

    def _configure_styles(self, document):
        self._configure_page_layout(document)

        for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
            style = document.styles[style_name]
            style.font.name = self.font_name

            if style._element.rPr is not None:
                style._element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)

        normal_style = document.styles["Normal"]
        normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15

        heading_configs = {
            "Heading 1": {"size": Pt(24), "bold": True, "space_after": Pt(12)},
            "Heading 2": {"size": Pt(18), "bold": True, "space_after": Pt(6)},
            "Heading 3": {"size": Pt(14), "bold": True, "space_after": Pt(6)},
            "Heading 4": {"size": Pt(12), "bold": True, "space_after": Pt(6)},
        }

        for heading_name, config in heading_configs.items():
            heading_style = document.styles[heading_name]
            heading_style.font.size = config["size"]
            heading_style.font.bold = config["bold"]
            heading_style.paragraph_format.space_after = config["space_after"]

            if heading_style._element.rPr is not None:
                heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)

    def _configure_page_layout(self, document):
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    def _set_run_font(self, run, font_name: Optional[str] = None, size: Optional[float] = None):
        final_font_name = font_name or self.font_name
        run.font.name = final_font_name

        if size:
            run.font.size = Pt(size)

        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), final_font_name)

    def _set_paragraph_font(self, paragraph, font_name: Optional[str] = None, size: Optional[float] = None):
        for run in paragraph.runs:
            self._set_run_font(run, font_name=font_name, size=size)

    def _add_content(self, document, content: str):
        lines = str(content or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
        cursor = 0
        in_code_block = False
        code_lines: List[str] = []

        while cursor < len(lines):
            line = lines[cursor]
            stripped = line.strip()

            if stripped.startswith("```"):

                if in_code_block:
                    self._add_code_block(document, code_lines)
                    in_code_block = False
                    code_lines = []
                else:
                    in_code_block = True
                    code_lines = []

                cursor += 1
                continue

            if in_code_block:
                code_lines.append(line)
                cursor += 1
                continue

            table_rows, next_cursor = self._read_table(lines, cursor)

            if table_rows:
                self._add_table(document, table_rows)
                cursor = next_cursor
                continue

            if not stripped:
                document.add_paragraph("")
                cursor += 1
                continue

            if self._is_horizontal_rule(stripped):
                self._add_horizontal_rule(document)
                cursor += 1
                continue

            if self._looks_like_html(stripped):
                self._add_html_line(document, line)
                cursor += 1
                continue

            if self._add_markdown_images(document, line):
                cursor += 1
                continue

            if self._add_markdown_block(document, line, stripped):
                cursor += 1
                continue

            self._add_inline_links(document, line.rstrip())
            cursor += 1

        if in_code_block:
            self._add_code_block(document, code_lines)

    def _add_markdown_block(self, document, line: str, stripped: str) -> bool:
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)

        if heading_match:
            level = min(len(heading_match.group(1)) + 1, 4)
            paragraph = document.add_heading(heading_match.group(2).strip(), level=level)
            self._set_paragraph_font(paragraph)
            return True

        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)

        if bullet_match:
            paragraph = document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            self._set_paragraph_font(paragraph)
            return True

        ordered_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)

        if ordered_match:
            paragraph = document.add_paragraph(ordered_match.group(1).strip(), style="List Number")
            self._set_paragraph_font(paragraph)
            return True

        quote_match = re.match(r"^>\s*(.+)$", stripped)

        if quote_match:
            self._add_blockquote(document, quote_match.group(1).strip())
            return True

        return False

    def _is_horizontal_rule(self, stripped: str) -> bool:
        compact = re.sub(r"\s+", "", str(stripped or ""))

        if len(compact) < 3:
            return False

        return bool(re.fullmatch(r"[-*_]+", compact) and len(set(compact)) == 1)

    def _add_horizontal_rule(self, document):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        p_pr = paragraph._element.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))

        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)

        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "6")
        bottom_border.set(qn("w:space"), "1")
        bottom_border.set(qn("w:color"), "BFBFBF")
        p_bdr.append(bottom_border)

    def _looks_like_html(self, line: str) -> bool:
        return bool(re.search(r"<\s*/?\s*(div|p|br|hr|img|span|h[1-6]|ul|ol|li|table|tr|td|th)\b", line, flags=re.I))

    def _add_html_line(self, document, line: str):
        parser = _HtmlContentParser()
        parser.feed(line)
        parser.close()

        text_parts: List[str] = []

        def flush_text():
            text = "".join(text_parts).strip()
            text_parts.clear()

            if text:
                paragraph = document.add_paragraph(text)
                self._set_paragraph_font(paragraph)

        for token in parser.tokens:
            token_type = token.get("type")

            if token_type == "text":
                text_parts.append(str(token.get("text") or ""))
                continue

            if token_type == "newline":
                flush_text()
                continue

            if token_type == "horizontal_rule":
                flush_text()
                self._add_horizontal_rule(document)
                continue

            if token_type == "image":
                flush_text()
                image_added = self._try_add_image(
                    document,
                    src=str(token.get("src") or ""),
                    alt=str(token.get("alt") or ""),
                    width_px=token.get("width_px"),
                    height_px=token.get("height_px"),
                )

                if not image_added:
                    self._add_plain_text(document, self._build_markdown_image_text(
                        src=str(token.get("src") or ""),
                        alt=str(token.get("alt") or ""),
                    ))

        flush_text()

    def _add_markdown_images(self, document, line: str) -> bool:
        matches = list(_MARKDOWN_IMAGE_RE.finditer(line))

        if not matches:
            return False

        cursor = 0

        for match in matches:
            leading_text = line[cursor:match.start()].strip()

            if leading_text:
                paragraph = document.add_paragraph(leading_text)
                self._set_paragraph_font(paragraph)

            image_added = self._try_add_image(document, src=match.group(2), alt=match.group(1))

            if not image_added:
                self._add_plain_text(document, match.group(0))

            cursor = match.end()

        trailing_text = line[cursor:].strip()

        if trailing_text:
            paragraph = document.add_paragraph(trailing_text)
            self._set_paragraph_font(paragraph)

        return True

    def _try_add_image(
        self,
        document,
        *,
        src: str,
        alt: str = "",
        width_px: Optional[float] = None,
        height_px: Optional[float] = None,
    ):
        try:
            image_source = self._resolve_image_source(src)
            kwargs = self._build_image_size_kwargs(width_px=width_px, height_px=height_px)
            paragraph = document.add_paragraph()
            paragraph.alignment = 1
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run()
            run.add_picture(image_source, **kwargs)
        except Exception as e:
            print(f"[FormatConvert] image export failed src={src} error={e}")
            return False

        if alt:
            caption = document.add_paragraph(str(alt).strip())
            caption.alignment = 1
            caption.paragraph_format.space_before = Pt(2)
            caption.paragraph_format.space_after = Pt(6)
            self._set_paragraph_font(caption, size=9)

        return True

    def _add_plain_text(self, document, text: str):
        paragraph = document.add_paragraph(str(text or ""))
        self._set_paragraph_font(paragraph)

    def _build_markdown_image_text(self, *, src: str, alt: str = "") -> str:
        return f"![{str(alt or '').strip()}]({str(src or '').strip()})"

    def _resolve_image_source(self, src: str):
        raw_src = str(src or "").strip()

        if not raw_src:
            raise ValueError("image src is empty")

        if raw_src.startswith("data:image/"):
            return BytesIO(self._decode_data_image(raw_src))

        parsed = urllib_parse.urlparse(raw_src)

        if parsed.path:
            local_path = self._resolve_local_knowledge_image(parsed.path)

            if local_path:
                return str(local_path)

        if parsed.scheme in ("http", "https"):
            return BytesIO(self._download_remote_image(raw_src))

        local_file = self._resolve_content_file(raw_src)

        if local_file:
            return str(local_file)

        raise ValueError(f"cannot resolve image src: {raw_src}")

    def _resolve_local_knowledge_image(self, url_path: str) -> Optional[Path]:
        match = _KNOWLEDGE_IMAGE_ROUTE_RE.match(str(url_path or "").strip())

        if not match:
            return None

        username = urllib_parse.unquote(match.group(1)).strip()
        image_id = urllib_parse.unquote(match.group(2)).strip().lower()

        if not _SAFE_USERNAME_RE.match(username) or not _SAFE_IMAGE_ID_RE.match(image_id):
            raise ValueError(f"invalid knowledge image path: {url_path}")

        image_root = (self.base_dir / "data" / "users" / username / "database" / "static" / "images").resolve()
        index_path = image_root / "index.json"

        if not index_path.exists():
            raise ValueError(f"knowledge image index missing: {url_path}")

        index_data = json.loads(index_path.read_text(encoding="utf-8-sig"))
        images = index_data.get("images") if isinstance(index_data, dict) else {}
        row = images.get(image_id) if isinstance(images, dict) else None

        if not isinstance(row, dict):
            raise ValueError(f"knowledge image missing: {url_path}")

        file_name = str(row.get("file_name") or "").strip()

        if not file_name:
            raise ValueError(f"knowledge image file name is empty: {url_path}")

        image_path = (image_root / file_name).resolve()

        if os.path.commonpath([str(image_root), str(image_path)]) != str(image_root):
            raise ValueError(f"knowledge image path escapes root: {url_path}")

        if not image_path.exists():
            raise ValueError(f"knowledge image file missing: {url_path}")

        return image_path

    def _resolve_content_file(self, src: str) -> Optional[Path]:
        raw_src = str(src or "").strip()

        if not raw_src:
            return None

        candidate = Path(raw_src)

        if not candidate.is_absolute():
            candidate = self.base_dir / raw_src.lstrip("/\\")

        resolved = candidate.resolve()

        try:
            if os.path.commonpath([str(self.base_dir), str(resolved)]) != str(self.base_dir):
                return None
        except ValueError:
            return None

        if resolved.exists() and resolved.is_file():
            return resolved

        return None

    def _download_remote_image(self, url: str) -> bytes:
        parsed = urllib_parse.urlparse(url)

        if parsed.scheme not in ("http", "https") or not parsed.netloc:
            raise ValueError(f"invalid remote image url: {url}")

        request = urllib_request.Request(url, headers={"User-Agent": "NexoraFormatConvert/1.0"})

        try:
            with urllib_request.urlopen(request, timeout=15) as response:
                raw = response.read(_IMAGE_MAX_BYTES + 1)
        except Exception as e:
            raise ValueError(f"remote image download failed: {url}") from e

        if len(raw) > _IMAGE_MAX_BYTES:
            raise ValueError(f"remote image too large: {url}")

        return raw

    def _decode_data_image(self, raw_src: str) -> bytes:
        header, sep, payload = raw_src.partition(",")

        if not sep or ";base64" not in header:
            raise ValueError("image data URL is not base64")

        try:
            raw = base64.b64decode(payload, validate=True)
        except Exception as e:
            raise ValueError("image data URL decode failed") from e

        if len(raw) > _IMAGE_MAX_BYTES:
            raise ValueError("image data URL too large")

        return raw

    def _build_image_size_kwargs(
        self,
        *,
        width_px: Optional[float] = None,
        height_px: Optional[float] = None,
    ) -> Dict[str, Any]:
        if width_px:
            width_inches = min(float(width_px) / 96.0, self.max_image_width_inches)
            return {"width": Inches(width_inches)}

        if height_px:
            height_inches = max(0.1, float(height_px) / 96.0)
            return {"height": Inches(height_inches)}

        return {"width": Inches(self.max_image_width_inches)}

    def _add_table(self, document, rows: List[List[str]]):
        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=0, cols=column_count)
        table.style = "Table Grid"

        for row_index, row in enumerate(rows):
            cells = table.add_row().cells

            for column_index in range(column_count):
                text = row[column_index] if column_index < len(row) else ""
                paragraph = cells[column_index].paragraphs[0]
                paragraph.text = ""
                run = paragraph.add_run(text)
                self._set_run_font(run)

                if row_index == 0:
                    run.bold = True
                    cell = cells[column_index]
                    shading = cell._element.get_or_add_tcPr()
                    shading.append(self._create_shading("D9D9D9"))

    def _create_shading(self, color: str):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        shading.set(qn("w:val"), "clear")
        return shading

    def _add_code_block(self, document, code_lines: List[str]):
        code_text = "\n".join(code_lines).rstrip()

        if not code_text:
            return

        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]

        shading_elm = self._create_shading("F5F5F5")
        tc_pr = cell._element.get_or_add_tcPr()
        tc_pr.append(shading_elm)
        tc_pr.set(qn("w:vAlign"), "top")

        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.left_indent = Pt(12)
        run = paragraph.add_run(code_text)
        self._set_run_font(run, font_name="Consolas", size=9)

    def _add_inline_links(self, document, text: str):
        link_re = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
        matches = list(link_re.finditer(text))

        if not matches:
            paragraph = document.add_paragraph(text)
            self._set_paragraph_font(paragraph)
            return

        paragraph = document.add_paragraph()
        self._set_paragraph_font(paragraph)
        last_end = 0

        for match in matches:
            link_text = match.group(1).strip()
            link_url = match.group(2).strip()

            if last_end < match.start():
                prefix = text[last_end:match.start()]

                if prefix:
                    run = paragraph.add_run(prefix)
                    self._set_run_font(run)

            if link_text and link_url:
                rel_id = document.part.relate_to(
                    link_url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )

                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), rel_id)
                run_elem = OxmlElement("w:r")
                r_pr = OxmlElement("w:rPr")
                r_fonts = OxmlElement("w:rFonts")
                r_fonts.set(qn("w:ascii"), self.font_name)
                r_fonts.set(qn("w:hAnsi"), self.font_name)
                r_fonts.set(qn("w:eastAsia"), self.font_name)
                r_pr.append(r_fonts)

                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "21")
                r_pr.append(sz)

                color = OxmlElement("w:color")
                color.set(qn("w:val"), "0000FF")
                r_pr.append(color)

                underline = OxmlElement("w:u")
                underline.set(qn("w:val"), "single")
                r_pr.append(underline)
                run_elem.append(r_pr)

                text_elem = OxmlElement("w:t")
                text_elem.text = link_text
                run_elem.append(text_elem)
                hyperlink.append(run_elem)
                paragraph._element.append(hyperlink)

            last_end = match.end()

        if last_end < len(text):
            suffix = text[last_end:]

            if suffix:
                run = paragraph.add_run(suffix)
                self._set_run_font(run)

    def _add_blockquote(self, document, text: str):
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]

        shading_elm = self._create_shading("F9F9F9")
        tc_pr = cell._element.get_or_add_tcPr()
        tc_pr.append(shading_elm)
        tc_pr.set(qn("w:vAlign"), "top")

        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.left_indent = Pt(24)

        run = paragraph.add_run(text)
        self._set_run_font(run)

        p_pr = paragraph._element.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        left_border = OxmlElement("w:left")
        left_border.set(qn("w:val"), "single")
        left_border.set(qn("w:sz"), "8")
        left_border.set(qn("w:space"), "4")
        left_border.set(qn("w:color"), "CCCCCC")
        p_bdr.append(left_border)
        p_pr.append(p_bdr)

    def _read_table(self, lines: List[str], start_index: int) -> Tuple[List[List[str]], int]:
        if start_index + 1 >= len(lines):
            return [], start_index

        header = self._split_table_row(lines[start_index])

        if not header or not self._is_table_separator(lines[start_index + 1]):
            return [], start_index

        rows = [header]
        cursor = start_index + 2

        while cursor < len(lines):
            row = self._split_table_row(lines[cursor])

            if not row:
                break

            rows.append(row)
            cursor += 1

        return rows, cursor

    def _split_table_row(self, line: str) -> List[str]:
        raw = str(line or "").strip()

        if "|" not in raw:
            return []

        if raw.startswith("|"):
            raw = raw[1:]

        if raw.endswith("|"):
            raw = raw[:-1]

        return [cell.strip() for cell in raw.split("|")]

    def _is_table_separator(self, line: str) -> bool:
        cells = self._split_table_row(line)

        if not cells:
            return False

        for cell in cells:
            marker = cell.replace(" ", "")

            if not marker or "-" not in marker:
                return False

            if marker.strip(":").replace("-", ""):
                return False

        return True
